"""Xuất báo cáo .docx "THÀNH QUẢ TÍNH TOÁN BÌNH SAI LƯỚI ĐỘ CAO" (SPEC.md mục 4).

Ghi chú triển khai: spec đề xuất điền vào file .docx mẫu có sẵn qua `docxtpl`
(placeholder + lặp bảng `{%tr for %}`). Chưa có file mẫu đã đánh dấu
placeholder được cung cấp, nên bản này dựng báo cáo trực tiếp bằng
`python-docx` theo đúng nội dung/thứ tự mục ở SPEC.md mục 4 (đối chiếu với
file mẫu Cầu Cửa Đại). Khi có file mẫu thật, có thể thay thế phần dựng
docx ở đây bằng `docxtpl` mà không cần đổi các hàm chuẩn bị dữ liệu
(`_build_context` có thể tái sử dụng làm context cho docxtpl).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from binh_sai_do_cao.core.accuracy import AccuracyExtremes, AccuracyResult
from binh_sai_do_cao.core.adjustment import AdjustmentResult
from binh_sai_do_cao.core.closure_check import ClosureResult
from binh_sai_do_cao.models.schema import NetworkInput

ADJUSTMENT_METHOD_LABELS = {"indirect": "Phụ thuộc", "conditional": "Điều kiện"}

DEFAULT_SOFTWARE_NAME = "binh_sai_do_cao"


def _fmt(value: float, decimals: int) -> str:
    return f"{value:.{decimals}f}"


def _fmt_signed(value: float, decimals: int) -> str:
    sign = "-" if value < 0 else ""
    return f"{sign}{abs(value):.{decimals}f}"


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.underline = True


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def build_report(
    data: NetworkInput,
    adjustment: AdjustmentResult,
    accuracy: AccuracyResult,
    extremes: AccuracyExtremes,
    closures: list[ClosureResult],
    software_name: str = DEFAULT_SOFTWARE_NAME,
) -> Document:
    doc = Document()

    # 1. Tiêu đề
    _add_title(doc, "THÀNH QUẢ TÍNH TOÁN BÌNH SAI LƯỚI ĐỘ CAO")

    # 2. Tên công trình
    doc.add_paragraph(f"Tên công trình: {data.project.name}")
    doc.add_paragraph()

    # 3. Chỉ tiêu kỹ thuật lưới
    _add_section_heading(doc, "Chỉ tiêu kỹ thuật lưới")
    k = data.settings.k_coefficient_mm
    n_var = "N" if data.settings.closure_formula == "k_sqrt_n" else "L"
    doc.add_paragraph(f"Số lượng điểm gốc\t\t: {len(data.known_points)}")
    doc.add_paragraph(f"Số lượng điểm mới lập\t : {len(adjustment.unknown_names)}")
    doc.add_paragraph(f"Số chênh cao đo\t\t : {len(data.observations)}")
    doc.add_paragraph(f"Sai số khép giới hạn\t\t : {_fmt(k, 1)} x SQRT({n_var}) mm")
    method_label = ADJUSTMENT_METHOD_LABELS.get(
        data.settings.adjustment_method, data.settings.adjustment_method
    )
    doc.add_paragraph(f"Phương pháp bình sai\t\t : {method_label}")
    doc.add_paragraph()

    # 4. Số liệu độ cao khởi tính
    _add_section_heading(doc, "Số liệu độ cao khởi tính")
    rows = [
        [str(i + 1), p.name, _fmt(p.height_m, 4), ""]
        for i, p in enumerate(data.known_points)
    ]
    _add_table(doc, ["TT", "Tên điểm", "H(m)", "Ghi chú"], rows)
    doc.add_paragraph()

    # 5. Bảng thành quả độ cao sau bình sai
    _add_section_heading(doc, "Bảng thành quả độ cao sau bình sai")
    rows = [
        [
            str(i + 1),
            name,
            _fmt(adjustment.heights_m[name], 4),
            _fmt(accuracy.m_point_mm[name], 2),
            "",
        ]
        for i, name in enumerate(adjustment.unknown_names)
    ]
    _add_table(doc, ["TT", "Tên điểm", "H(m)", "Sai số TP (mm)", "Ghi chú"], rows)
    doc.add_paragraph()

    # 6. Bảng trị đo, SHC, trị bình sai chênh cao
    _add_section_heading(doc, "Bảng trị đo, số hiệu chỉnh trị bình sai chênh cao")
    rows = [
        [
            str(i + 1),
            obs.from_,
            obs.to,
            _fmt_signed(obs.measured_dh_m, 4),
            _fmt_signed(adjustment.corrections_m[obs.id], 4),
            _fmt_signed(adjustment.adjusted_dh_m[obs.id], 4),
            _fmt(accuracy.m_obs_mm[obs.id], 2),
            str(obs.stations if obs.stations is not None else ""),
        ]
        for i, obs in enumerate(data.observations)
    ]
    _add_table(
        doc,
        ["TT", "Điểm đầu", "Điểm cuối", "Trị đo (m)", "SHC (m)", "Trị BS (m)", "SSTP (mm)", "Trạm đo"],
        rows,
    )
    doc.add_paragraph()

    # 7. Kết quả đánh giá độ chính xác lưới
    _add_section_heading(doc, "Kết quả đánh giá độ chính xác lưới")
    doc.add_paragraph(f"1. Sai số trung phương trọng số đơn vị\t: Mo = {_fmt(accuracy.mo_mm, 2)}(mm/Tr)")
    doc.add_paragraph(
        f"2. Sai số trung phương Độ cao lớn nhất\t: ({extremes.point_max_name}) = "
        f"{_fmt(extremes.point_max_mm, 2)}(mm)"
    )
    doc.add_paragraph(
        f"3. Sai số trung phương Độ cao nhỏ nhất\t: ({extremes.point_min_name}) = "
        f"{_fmt(extremes.point_min_mm, 2)}(mm)"
    )
    obs_max_from, obs_max_to = extremes.obs_max_pair
    obs_min_from, obs_min_to = extremes.obs_min_pair
    doc.add_paragraph(
        f"4. Sai số trị đo chênh cao lớn nhất\t\t: ({obs_max_from} → {obs_max_to}) = "
        f"{_fmt(extremes.obs_max_mm, 2)}(mm)"
    )
    doc.add_paragraph(
        f"5. Sai số trị đo chênh cao nhỏ nhất\t\t: ({obs_min_from} → {obs_min_to}) = "
        f"{_fmt(extremes.obs_min_mm, 2)}(mm)"
    )
    doc.add_paragraph()

    # 8. Kết quả kiểm tra sai số khép
    _add_section_heading(doc, "Kết quả kiểm tra sai số khép")
    for i, c in enumerate(closures, start=1):
        suffix = f" ({c.flag})" if c.flag else ""
        doc.add_paragraph(f"{i}. Tuyến: {c.loop.path_str}")
        doc.add_paragraph(f"a. Số đoạn đo\t\t:  N\t\t= {c.loop.n_segments}")
        doc.add_paragraph(f"b. Tổng số trạm đo\t\t:  [N]\t\t= {c.loop.total_stations}")
        doc.add_paragraph(f"c. Sai số khép độ cao\t:  Wh\t\t= {_fmt_signed(c.wh_mm, 2)} (mm)")
        doc.add_paragraph(
            f"d. Sai số khép giới hạn\t:  Wh(gh)\t= ±{_fmt(c.wh_gh_mm, 2)} (mm){suffix}"
        )
        if i < len(closures):
            doc.add_paragraph("_" * 79)
    doc.add_paragraph()

    # 9. Chữ ký
    date_line = f"Ngày {data.project.date}" if data.project.date else "Ngày ... tháng ... năm ..."
    date_p = doc.add_paragraph(date_line)
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Người đo đạc   : {data.project.surveyor}")
    doc.add_paragraph(f"Người tính toán: {data.project.computer}")
    doc.add_paragraph(f"Người kiểm tra : {data.project.checker}")

    # 10. Footer
    doc.add_paragraph(f"Kết quả được tính toán bằng phần mềm {software_name}")

    return doc


def write_report(
    path: str | Path,
    data: NetworkInput,
    adjustment: AdjustmentResult,
    accuracy: AccuracyResult,
    extremes: AccuracyExtremes,
    closures: list[ClosureResult],
    software_name: str = DEFAULT_SOFTWARE_NAME,
) -> None:
    doc = build_report(data, adjustment, accuracy, extremes, closures, software_name)
    doc.save(str(path))
